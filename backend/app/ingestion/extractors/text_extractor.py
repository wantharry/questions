"""
Text and other document format extractors.
"""
from pathlib import Path
from typing import Dict, Any
from docx import Document as DocxDocument
import markdown
from app.utils.logger import app_logger


class TextExtractor:
    """Extract content from plain text files."""
    
    def extract(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            metadata = {
                "file_size": file_path.stat().st_size,
            }
            
            app_logger.info(f"Extracted {len(text)} characters from text file {file_path.name}")
            
            return {
                "text": text,
                "metadata": metadata,
            }
        
        except Exception as e:
            app_logger.error(f"Error extracting text file {file_path}: {e}")
            raise


class MarkdownExtractor:
    """Extract content from Markdown files."""
    
    def extract(self, file_path: Path) -> Dict[str, Any]:
        """Extract and convert Markdown to plain text."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = f.read()
            
            # Convert to HTML then extract text (preserves structure better)
            html = markdown.markdown(md_content)
            
            # Simple HTML tag removal for plain text
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            
            metadata = {
                "file_size": file_path.stat().st_size,
                "original_format": "markdown",
            }
            
            app_logger.info(f"Extracted {len(text)} characters from Markdown {file_path.name}")
            
            return {
                "text": text,
                "metadata": metadata,
            }
        
        except Exception as e:
            app_logger.error(f"Error extracting Markdown {file_path}: {e}")
            raise


class DocxExtractor:
    """Extract content from Word documents (.docx)."""
    
    def extract(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = '\n'.join(paragraphs)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    tables_text.append(row_text)
            
            if tables_text:
                text += '\n\nTables:\n' + '\n'.join(tables_text)
            
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "file_size": file_path.stat().st_size,
            }
            
            # Try to get core properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.subject:
                    metadata["subject"] = core_props.subject
            except:
                pass
            
            app_logger.info(f"Extracted {len(text)} characters from DOCX {file_path.name}")
            
            return {
                "text": text,
                "metadata": metadata,
            }
        
        except Exception as e:
            app_logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
